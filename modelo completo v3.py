from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
from dotenv import load_dotenv
from supabase import create_client
from urllib.parse import quote
from datetime import datetime
import os

# Carrega variáveis do ambiente
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
BUCKET = "peticoesgeradas"

# Inicializa clientes
client = OpenAI(api_key=OPENAI_API_KEY)
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

# Inicializa API
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Conta tokens simples
def contar_tokens(texto):
    return len(texto.split())

# Gera nome de arquivo único com reclamante + timestamp
def gerar_nome_arquivo(dados_usuario):
    nome_base = dados_usuario.get('reclamante', 'anonimo').replace(' ', '_')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"peticao_{nome_base}_{timestamp}.docx"

# IA reescreve parágrafos com base no contexto e seção
def reescrever_com_ia(paragrafos, contexto_usuario: dict):
    texto_final = ""
    custo_total = 0
    contexto = "\n".join([f"{chave}: {valor}" for chave, valor in contexto_usuario.items()])
    secao_atual = ""

    for par in paragrafos:
        if not par.strip():
            texto_final += "\n"
            continue

        texto = par.strip()

        if texto.isupper() and len(texto) < 100:
            secao_atual = texto
            texto_final += texto + "\n\n"
            continue

        modelo = "gpt-4" if any(s in secao_atual for s in ["FATOS", "FUNDAMENTAÇÃO", "PEDIDOS"]) else "gpt-3.5-turbo"

        prompt = (
            f"Você é um advogado redigindo uma petição trabalhista.\n"
            f"Abaixo estão os dados fornecidos pelo cliente:\n\n{contexto}\n\n"
            f"Seção: {secao_atual or 'Geral'}\n\n"
            f"Reescreva o seguinte parágrafo com linguagem jurídica clara, formal e adaptada:\n\n{par}"
        )

        try:
            resposta = client.chat.completions.create(
                model=modelo,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.4
            )
            texto_revisado = resposta.choices[0].message.content.strip()
            print(f"\n--- [{modelo}] Seção: {secao_atual} ---\n{texto}\n→ {texto_revisado}\n")
        except Exception as e:
            texto_revisado = f"[ERRO NA IA] {par}"
            print(f"Erro: {e}")

        texto_final += texto_revisado + "\n\n"

        tokens_in = contar_tokens(prompt)
        tokens_out = contar_tokens(texto_revisado)
        custo = (tokens_in / 1000) * (0.01 if modelo == "gpt-4" else 0.0005) + \
                (tokens_out / 1000) * (0.03 if modelo == "gpt-4" else 0.0015)
        custo_total += custo

    return texto_final.strip(), custo_total

# Formata texto no Word com estilo jurídico
def formatar_documento_visualmente(texto_formatado, caminho_saida, caminho_imagem_cabecalho=None):
    doc = Document()

    if caminho_imagem_cabecalho and os.path.isfile(caminho_imagem_cabecalho):
        section = doc.sections[0]
        header = section.header
        for p in header.paragraphs:
            p.clear()
        p = header.paragraphs[0]
        run = p.add_run()
        run.add_picture(caminho_imagem_cabecalho, width=Inches(2.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for par in texto_formatado.split("\n"):
        texto = par.strip()
        if not texto:
            doc.add_paragraph("")
            continue
        if texto.isupper() and len(texto) < 100:
            p = doc.add_paragraph()
            run = p.add_run(texto)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph()
            run = p.add_run(texto)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "IBOTI ADVOCACIA – OAB/RS 65.382 | www.ibotiadvogados.com.br | (51) 98418-2882"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(caminho_saida)

# Envia o .docx para o Supabase e retorna a URL pública
def enviar_para_supabase(caminho_local, nome_arquivo):
    with open(caminho_local, "rb") as f:
        dados = f.read()

    caminho_remoto = nome_arquivo
    supabase.storage.from_(BUCKET).upload(caminho_remoto, dados, {
        "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    })

    url = f"{SUPABASE_URL}/storage/v1/object/public/{BUCKET}/{quote(nome_arquivo)}"
    return url

# Endpoint principal
@app.post("/gerar-peticao")
async def gerar_peticao(request: Request):
    dados = await request.json()

    nome_arquivo = gerar_nome_arquivo(dados)
    caminho_modelo = "c:/Users/escri/Documents/Petição inicial/modelos/modelo_base.docx"
    caminho_saida = f"./output/{nome_arquivo}"
    imagem_cabecalho = "C:/Users/escri/Documents/logo_iboti.jpg"

    os.makedirs("output", exist_ok=True)

    if not os.path.isfile(caminho_modelo):
        return {"erro": f"Modelo não encontrado: {caminho_modelo}"}

    doc = Document(caminho_modelo)
    paragrafos = [p.text for p in doc.paragraphs]

    texto_revisado, custo = reescrever_com_ia(paragrafos, dados)
    formatar_documento_visualmente(texto_revisado, caminho_saida, imagem_cabecalho)
    url_download = enviar_para_supabase(caminho_saida, nome_arquivo)

    return {
        "arquivo": nome_arquivo,
        "url_download": url_download,
        "custo_estimado_usd": round(custo, 4)
    }

# Rota preflight para CORS
@app.options("/gerar-peticao")
async def options_gerar_peticao():
    return {"status": "ok"}
